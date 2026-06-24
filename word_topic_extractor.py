import argparse
import re
from pathlib import Path

from docx import Document


HEADING_RE = re.compile(r"^Heading\s*(\d+)$", re.IGNORECASE)
TOC_RE = re.compile(r"^TOC\s*(\d+)$", re.IGNORECASE)
TRAILING_PAGE_RE = re.compile(r"^(?P<title>.+?)[\s.\t]+(?P<page>\d+)$")


def heading_level(paragraph):
    match = HEADING_RE.match(paragraph.style.name or "")
    if not match:
        return None
    return int(match.group(1))


def normalize(text):
    return " ".join(text.casefold().split())


def iter_body_items(document):
    body = document.element.body
    for element in body:
        if element.tag.endswith("}p"):
            yield "paragraph", element
        elif element.tag.endswith("}tbl"):
            yield "table", element


def paragraph_text(element):
    return "".join(node.text or "" for node in element.iter() if node.tag.endswith("}t")).strip()


def clean_toc_text(text):
    text = re.sub(r"\s+", " ", text.replace("\t", " ")).strip()
    return re.sub(r"\.{2,}", " ", text).strip()


def paragraph_style_name(element):
    p_style = element.find(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pStyle")
    if p_style is None:
        return ""
    return p_style.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "")


def element_heading_level(element):
    style_name = paragraph_style_name(element).replace("_", " ")
    match = HEADING_RE.match(style_name)
    if not match:
        return None
    return int(match.group(1))


def find_topic_range(document, topic):
    wanted = normalize(topic)
    items = list(iter_body_items(document))

    start_index = None
    start_level = None

    for index, (kind, element) in enumerate(items):
        if kind != "paragraph":
            continue

        text = paragraph_text(element)
        if normalize(text) == wanted:
            start_index = index
            start_level = element_heading_level(element) or 1
            break

    if start_index is None:
        for index, (kind, element) in enumerate(items):
            if kind != "paragraph":
                continue

            text = normalize(paragraph_text(element))
            if wanted and wanted in text:
                start_index = index
                start_level = element_heading_level(element) or 1
                break

    if start_index is None:
        return None

    end_index = len(items)
    for index in range(start_index + 1, len(items)):
        kind, element = items[index]
        if kind != "paragraph":
            continue

        level = element_heading_level(element)
        if level is not None and level <= start_level:
            end_index = index
            break

    return start_index, end_index


def list_topics(document):
    topics = []
    for paragraph in document.paragraphs:
        level = heading_level(paragraph)
        text = paragraph.text.strip()
        if normalize(text) in {"contents", "table of contents"}:
            continue
        if level and text:
            topics.append((level, text))
    return topics


def toc_level(paragraph):
    style_name = paragraph.style.name or ""
    match = TOC_RE.match(style_name)
    if not match:
        return None
    return int(match.group(1))


def parse_toc_entry(paragraph):
    text = clean_toc_text(paragraph.text)
    if not text:
        return None

    match = TRAILING_PAGE_RE.match(text)
    if match:
        title = match.group("title").strip()
        page = match.group("page").strip()
    else:
        title = text.strip()
        page = ""

    if not title:
        return None

    return {
        "level": toc_level(paragraph) or 1,
        "title": title,
        "page": page,
        "source": "toc",
    }


def plain_contents_entries(document):
    entries = []
    in_contents = False

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        level = heading_level(paragraph)

        if normalize(text) in {"contents", "table of contents"}:
            in_contents = True
            continue

        if not in_contents:
            continue

        if level is not None:
            break

        entry = parse_toc_entry(paragraph)
        if entry and entry["page"]:
            entries.append(entry)

    return entries


def content_table_entries(document):
    toc_entries = []
    for paragraph in document.paragraphs:
        if toc_level(paragraph):
            entry = parse_toc_entry(paragraph)
            if entry:
                toc_entries.append(entry)

    if toc_entries:
        return toc_entries

    plain_entries = plain_contents_entries(document)
    if plain_entries:
        return plain_entries

    return [
        {"level": level, "title": title, "page": "", "source": "heading"}
        for level, title in list_topics(document)
    ]


def copy_core_properties(source, target):
    target.core_properties.author = source.core_properties.author
    target.core_properties.category = source.core_properties.category
    target.core_properties.comments = source.core_properties.comments
    target.core_properties.keywords = source.core_properties.keywords
    target.core_properties.subject = source.core_properties.subject
    target.core_properties.title = source.core_properties.title


def selected_topic_ranges(source, selected_topics):
    ranges = []
    missing = []
    for topic in selected_topics:
        topic_range = find_topic_range(source, topic)
        if topic_range is None:
            missing.append(topic)
        else:
            ranges.append((topic_range[0], topic_range[1], topic))

    if missing:
        available = ", ".join(title for _, title in list_topics(source)) or "no styled headings found"
        raise ValueError(f"Topic not found: {', '.join(missing)}. Available topics: {available}")

    ranges.sort(key=lambda item: item[0])
    return merge_ranges(ranges)


def merge_ranges(ranges):
    merged = []
    for start, end, topic in ranges:
        if merged and start < merged[-1][1]:
            continue
        merged.append((start, end, topic))
    return merged


def keep_only_ranges(document, ranges):
    keep_indexes = set()
    for start, end, _ in ranges:
        keep_indexes.update(range(start, end))

    body = document.element.body
    items = list(iter_body_items(document))
    keep_elements = {element for index, (_, element) in enumerate(items) if index in keep_indexes}

    for child in list(body):
        if child.tag.endswith("}sectPr"):
            continue
        if child not in keep_elements:
            body.remove(child)


def extract_topic(input_path, topic, output_path):
    return extract_topics(input_path, [topic], output_path)


def extract_topics(input_path, topics, output_path):
    selected_topics = [topic for topic in topics if topic and topic.strip()]
    if not selected_topics:
        raise ValueError("Please select at least one topic.")

    output = Document(input_path)
    ranges = selected_topic_ranges(output, selected_topics)
    keep_only_ranges(output, ranges)

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output.save(output_path)
    return output_path


def build_parser():
    parser = argparse.ArgumentParser(
        description="Extract one selected topic from a complete Word document into a new Word document."
    )
    parser.add_argument("input", help="Path to the complete .docx file.")
    parser.add_argument("topic", nargs="?", help="Topic or heading to extract.")
    parser.add_argument("-o", "--output", default=None, help="Output .docx path.")
    parser.add_argument("--list-topics", action="store_true", help="Show headings found in the document.")
    return parser


def main():
    parser = build_parser()
    args = parser.parse_args()

    input_path = Path(args.input)
    if not input_path.exists():
        parser.error(f"Input file does not exist: {input_path}")

    document = Document(input_path)

    if args.list_topics:
        topics = content_table_entries(document)
        if not topics:
            print("No table-of-contents entries or heading-style topics found.")
            return

        print("Available topics:")
        for entry in topics:
            indent = "  " * (entry["level"] - 1)
            page = f" (page {entry['page']})" if entry["page"] else ""
            print(f"{indent}- {entry['title']}{page}")
        return

    if not args.topic:
        parser.error("Please provide a topic, or use --list-topics.")

    safe_topic = re.sub(r"[^A-Za-z0-9._-]+", "_", args.topic.strip()).strip("_") or "selected_topic"
    output_path = args.output or f"outputs/{safe_topic}.docx"
    result = extract_topic(input_path, args.topic, output_path)
    print(f"Created {result}")


if __name__ == "__main__":
    main()
