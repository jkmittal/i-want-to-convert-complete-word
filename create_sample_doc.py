from docx import Document


def add_topic(document, title, paragraphs, subtopics=None):
    document.add_heading(title, level=1)
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)

    for subtitle, lines in (subtopics or {}).items():
        document.add_heading(subtitle, level=2)
        for line in lines:
            document.add_paragraph(line)


def main():
    document = Document()
    document.add_heading("Complete Technology Notes", level=0)
    document.add_paragraph(
        "This sample file contains multiple topics. Use the extractor to create a new "
        "Word file for only one selected topic."
    )

    document.add_heading("Contents", level=1)
    contents = [
        ("Artificial Intelligence", "1", "TOC 1"),
        ("Common Uses", "2", "TOC 2"),
        ("Machine Learning", "3", "TOC 1"),
        ("Types of Machine Learning", "4", "TOC 2"),
        ("Applications", "5", "TOC 2"),
        ("Cyber Security", "6", "TOC 1"),
    ]
    for title, page, style in contents:
        paragraph = document.add_paragraph()
        try:
            paragraph.style = style
        except KeyError:
            pass
        paragraph.add_run(f"{title}\t{page}")

    add_topic(
        document,
        "Artificial Intelligence",
        [
            "Artificial intelligence is the field of building systems that can perform tasks requiring human-like reasoning.",
            "It includes planning, learning, perception, language understanding, and decision making.",
        ],
        {
            "Common Uses": [
                "AI is used in chatbots, recommendation systems, fraud detection, and automation.",
                "Modern AI systems often combine data, algorithms, and feedback loops.",
            ]
        },
    )

    add_topic(
        document,
        "Machine Learning",
        [
            "Machine learning is a branch of AI where systems learn patterns from data instead of being explicitly programmed for every rule.",
            "A model is trained on examples and then used to make predictions or decisions.",
        ],
        {
            "Types of Machine Learning": [
                "Supervised learning uses labeled examples.",
                "Unsupervised learning finds structure in unlabeled data.",
                "Reinforcement learning improves decisions through rewards and feedback.",
            ],
            "Applications": [
                "Machine learning is used for image recognition, forecasting, medical screening, and personalization.",
            ],
        },
    )

    add_topic(
        document,
        "Cyber Security",
        [
            "Cyber security protects systems, networks, and data from unauthorized access or damage.",
            "Important practices include strong passwords, software updates, access control, and monitoring.",
        ],
    )

    document.save("sample_complete_document.docx")
    print("Created sample_complete_document.docx")


if __name__ == "__main__":
    main()
